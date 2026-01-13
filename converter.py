#!/usr/bin/env python3
"""
MD2CV - Markdown to CV Converter

Converts Markdown CVs to ATS-compatible PDF and DOCX formats.
Supports custom templates and YAML frontmatter for metadata.

Usage:
    python converter.py input.md --format pdf --template ats_classic
    python converter.py input.md --format docx --template ats_classic
    python converter.py input.md --format all --template ats_classic
"""

import argparse
import os
import re
import sys
from pathlib import Path
from typing import Optional

import markdown
import yaml
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


# --- Constants ---
SCRIPT_DIR = Path(__file__).parent.resolve()
TEMPLATES_DIR = SCRIPT_DIR / "templates"
OUTPUT_DIR = SCRIPT_DIR / "output"


# --- Markdown Parser ---
class MarkdownParser:
    """Parse Markdown content with YAML frontmatter and HTML support."""
    
    FRONTMATTER_PATTERN = re.compile(r'^---\s*\n(.*?)\n---\s*\n', re.DOTALL)
    
    def __init__(self, content: str):
        self.raw_content = content
        self.meta: dict = {}
        self.body: str = ""
        self._parse()
    
    def _parse(self):
        """Extract YAML frontmatter and markdown body."""
        match = self.FRONTMATTER_PATTERN.match(self.raw_content)
        
        if match:
            yaml_content = match.group(1)
            self.meta = yaml.safe_load(yaml_content) or {}
            self.body = self.raw_content[match.end():]
        else:
            self.meta = {}
            self.body = self.raw_content
    
    def to_html(self) -> str:
        """Convert markdown body to HTML."""
        md = markdown.Markdown(
            extensions=[
                'extra',      # Tables, fenced code, footnotes, etc.
                'meta',       # Meta-data
                'nl2br',      # Newlines to <br>
                'sane_lists', # Better list handling
                'smarty',     # Smart quotes
            ]
        )
        return md.convert(self.body)


# --- Template Renderer ---
class TemplateRenderer:
    """Render CV content using Jinja2 templates."""
    
    def __init__(self, template_name: str):
        self.template_name = template_name
        self.template_dir = TEMPLATES_DIR / template_name
        
        if not self.template_dir.exists():
            available = self._list_templates()
            raise ValueError(
                f"Template '{template_name}' not found.\n"
                f"Available templates: {', '.join(available)}"
            )
        
        self.env = Environment(
            loader=FileSystemLoader(str(self.template_dir)),
            autoescape=True
        )
    
    def _list_templates(self) -> list[str]:
        """List available template names."""
        if not TEMPLATES_DIR.exists():
            return []
        return [d.name for d in TEMPLATES_DIR.iterdir() if d.is_dir()]
    
    def render(self, meta: dict, content: str) -> str:
        """Render the template with metadata and content."""
        template = self.env.get_template("template.html")
        return template.render(meta=meta, content=content)
    
    def get_css_path(self) -> Path:
        """Get the path to the template's CSS file."""
        return self.template_dir / "style.css"


# --- PDF Exporter ---
class PDFExporter:
    """Export rendered HTML to PDF using WeasyPrint."""
    
    def __init__(self, renderer: TemplateRenderer):
        self.renderer = renderer
    
    def export(self, html_content: str, output_path: Path):
        """Generate PDF from HTML content."""
        css_path = self.renderer.get_css_path()
        
        # Create HTML document
        html = HTML(string=html_content, base_url=str(self.renderer.template_dir))
        
        # Load CSS if exists
        stylesheets = []
        if css_path.exists():
            stylesheets.append(CSS(filename=str(css_path)))
        
        # Generate PDF
        html.write_pdf(str(output_path), stylesheets=stylesheets)
        print(f"✓ PDF created: {output_path}")


# --- DOCX Exporter ---
class DOCXExporter:
    """Export CV content to DOCX format."""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Configure document styles for ATS compatibility."""
        styles = self.doc.styles
        
        # Normal style
        normal = styles['Normal']
        normal.font.name = 'Georgia'
        normal.font.size = Pt(11)
        
        # Heading 1 - Section headers
        h1 = styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(14)
        h1.font.bold = True
        
        # Heading 2 - Subsection headers
        h2 = styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(12)
        h2.font.bold = True
    
    def export(self, meta: dict, markdown_body: str, output_path: Path):
        """Generate DOCX from metadata and markdown content."""
        # Add header with name
        if meta.get('name'):
            title = self.doc.add_heading(meta['name'], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add title/position
        if meta.get('title'):
            subtitle = self.doc.add_paragraph(meta['title'])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].italic = True
        
        # Add contact info
        contact_parts = []
        if meta.get('email'):
            contact_parts.append(meta['email'])
        if meta.get('phone'):
            contact_parts.append(meta['phone'])
        if meta.get('location'):
            contact_parts.append(meta['location'])
        
        if contact_parts:
            contact = self.doc.add_paragraph(' | '.join(contact_parts))
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add links
        link_parts = []
        if meta.get('linkedin'):
            link_parts.append(meta['linkedin'])
        if meta.get('github'):
            link_parts.append(meta['github'])
        if meta.get('website'):
            link_parts.append(meta['website'])
        
        if link_parts:
            links = self.doc.add_paragraph(' | '.join(link_parts))
            links.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add horizontal line (via paragraph border)
        self.doc.add_paragraph('_' * 60)
        
        # Parse and add markdown content
        self._add_markdown_content(markdown_body)
        
        # Save document
        self.doc.save(str(output_path))
        print(f"✓ DOCX created: {output_path}")
    
    def _add_markdown_content(self, markdown_body: str):
        """Parse markdown and add to document."""
        lines = markdown_body.strip().split('\n')
        current_list = []
        
        for line in lines:
            stripped = line.strip()
            
            # Skip empty lines
            if not stripped:
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                continue
            
            # Skip horizontal rules and HTML comments
            if stripped.startswith('---') or stripped.startswith('<!--'):
                continue
            
            # Heading 1
            if stripped.startswith('# '):
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                self.doc.add_heading(stripped[2:], level=1)
            
            # Heading 2
            elif stripped.startswith('## '):
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                self.doc.add_heading(stripped[3:], level=2)
            
            # Heading 3
            elif stripped.startswith('### '):
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                p = self.doc.add_paragraph(stripped[4:])
                p.runs[0].italic = True
            
            # List items
            elif stripped.startswith('- ') or stripped.startswith('* '):
                current_list.append(stripped[2:])
            
            # Bold text (standalone line)
            elif stripped.startswith('**') and stripped.endswith('**'):
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                text = stripped[2:-2]
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
            
            # Regular paragraph
            else:
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                # Clean up markdown formatting
                text = self._clean_markdown(stripped)
                if text:
                    self.doc.add_paragraph(text)
        
        # Add remaining list items
        if current_list:
            self._add_list(current_list)
    
    def _add_list(self, items: list[str]):
        """Add bullet list to document."""
        for item in items:
            text = self._clean_markdown(item)
            p = self.doc.add_paragraph(text, style='List Bullet')
    
    def _clean_markdown(self, text: str) -> str:
        """Remove markdown formatting from text."""
        # Remove bold markers
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        # Remove italic markers
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        # Remove links but keep text
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        return text.strip()


# --- Main Converter ---
class CVConverter:
    """Main converter class orchestrating the conversion process."""
    
    def __init__(self, input_path: str, template_name: str = "ats_classic"):
        self.input_path = Path(input_path)
        self.template_name = template_name
        
        if not self.input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        # Ensure output directory exists
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    def convert(self, formats: list[str]) -> list[Path]:
        """Convert markdown to specified formats."""
        # Read and parse markdown
        content = self.input_path.read_text(encoding='utf-8')
        parser = MarkdownParser(content)
        
        # Render HTML
        renderer = TemplateRenderer(self.template_name)
        html_content = renderer.render(parser.meta, parser.to_html())
        
        # Determine output base name
        output_base = OUTPUT_DIR / self.input_path.stem
        created_files = []
        
        # Export to requested formats
        if 'pdf' in formats or 'all' in formats:
            pdf_path = output_base.with_suffix('.pdf')
            pdf_exporter = PDFExporter(renderer)
            pdf_exporter.export(html_content, pdf_path)
            created_files.append(pdf_path)
        
        if 'docx' in formats or 'all' in formats:
            docx_path = output_base.with_suffix('.docx')
            docx_exporter = DOCXExporter()
            docx_exporter.export(parser.meta, parser.body, docx_path)
            created_files.append(docx_path)
        
        return created_files


# --- CLI ---
def list_templates():
    """List available templates."""
    print("\nAvailable templates:")
    print("-" * 40)
    
    if not TEMPLATES_DIR.exists():
        print("  No templates directory found.")
        return
    
    for template_dir in TEMPLATES_DIR.iterdir():
        if template_dir.is_dir():
            has_html = (template_dir / "template.html").exists()
            has_css = (template_dir / "style.css").exists()
            status = "✓" if has_html else "✗"
            print(f"  {status} {template_dir.name}")
            
            if not has_html:
                print(f"      └─ Missing: template.html")


def main():
    """Main entry point for CLI."""
    parser = argparse.ArgumentParser(
        description="MD2CV - Convert Markdown CVs to PDF and DOCX",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python converter.py cv.md --format pdf
  python converter.py cv.md --format docx --template ats_classic
  python converter.py cv.md --format all
  python converter.py --list-templates
        """
    )
    
    parser.add_argument(
        "input",
        nargs="?",
        help="Input Markdown file path"
    )
    
    parser.add_argument(
        "-f", "--format",
        choices=["pdf", "docx", "all"],
        default="all",
        help="Output format (default: all)"
    )
    
    parser.add_argument(
        "-t", "--template",
        default="ats_classic",
        help="Template name (default: ats_classic)"
    )
    
    parser.add_argument(
        "-o", "--output-dir",
        help="Output directory (default: ./output)"
    )
    
    parser.add_argument(
        "--list-templates",
        action="store_true",
        help="List available templates"
    )
    
    args = parser.parse_args()
    
    # Handle list templates
    if args.list_templates:
        list_templates()
        return
    
    # Validate input
    if not args.input:
        parser.error("Input file is required (unless using --list-templates)")
    
    # Override output directory if specified
    global OUTPUT_DIR
    if args.output_dir:
        OUTPUT_DIR = Path(args.output_dir)
    
    try:
        # Run conversion
        converter = CVConverter(args.input, args.template)
        formats = [args.format] if args.format != "all" else ["pdf", "docx"]
        created_files = converter.convert(formats)
        
        print(f"\n✓ Conversion complete! {len(created_files)} file(s) created.")
        
    except Exception as e:
        print(f"\n✗ Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
